import csv
from io import BytesIO
import json
from pathlib import Path

from docx import Document
from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload

from app.config import get_settings
from app.models import Company, ReportArtifact, ResearchResult, ResearchTask


def _result_rows(session: Session, task: ResearchTask) -> list[ResearchResult]:
    stmt = (
        select(ResearchResult)
        .where(ResearchResult.task_id == task.id)
        .options(
            selectinload(ResearchResult.company).selectinload(Company.sources),
            selectinload(ResearchResult.company).selectinload(Company.contacts),
        )
        .order_by(ResearchResult.rank.asc().nullslast(), ResearchResult.relevance_score.desc())
    )
    return list(session.scalars(stmt).all())


def _report_dir(task_id: int) -> Path:
    settings = get_settings()
    directory = Path(settings.report_output_dir) / f"task-{task_id}"
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def export_task_reports(session: Session, task: ResearchTask) -> dict[str, str]:
    rows = _result_rows(session, task)
    directory = _report_dir(task.id)
    paths = {
        "json": str(directory / "report.json"),
        "csv": str(directory / "report.csv"),
        "markdown": str(directory / "report.md"),
    }
    _write_json(paths["json"], task, rows)
    _write_csv(paths["csv"], rows)
    _write_markdown(paths["markdown"], task, rows)
    _store_docx(session, task, rows)
    return paths


def _write_json(path: str, task: ResearchTask, rows: list[ResearchResult]) -> None:
    payload = {
        "task": {
            "id": task.id,
            "target_industry": task.target_industry,
            "target_geography": task.target_geography,
            "si_keywords": task.si_keywords,
            "exclusion_keywords": task.exclusion_keywords,
            "service_categories": task.service_categories,
            "max_results": task.max_results,
            "search_queries": task.search_queries,
        },
        "caveats": [
            "Starter search provider uses deterministic public placeholders until a real search API is configured.",
            "Only public business contact fields should be collected by future providers.",
        ],
        "results": [_serialize_result(row) for row in rows],
    }
    Path(path).write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _write_csv(path: str, rows: list[ResearchResult]) -> None:
    with Path(path).open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "rank",
                "company_name",
                "website",
                "headquarters",
                "score",
                "services",
                "vendor_partnerships",
                "sources",
                "notes",
            ],
        )
        writer.writeheader()
        for row in rows:
            company = row.company
            writer.writerow(
                {
                    "rank": row.rank,
                    "company_name": company.name,
                    "website": company.website,
                    "headquarters": company.headquarters,
                    "score": row.relevance_score,
                    "services": "; ".join(company.services),
                    "vendor_partnerships": "; ".join(company.vendor_partnerships),
                    "sources": "; ".join(source.source_url for source in company.sources),
                    "notes": row.notes or company.notes,
                }
            )


def _write_markdown(path: str, task: ResearchTask, rows: list[ResearchResult]) -> None:
    lines = [
        f"# SI Research Report: Task {task.id}",
        "",
        "## Task Summary",
        "",
        f"- Target industry: {task.target_industry}",
        f"- Target geography: {task.target_geography}",
        f"- Services: {', '.join(task.service_categories) or 'Any'}",
        f"- Maximum results: {task.max_results}",
        "",
        "## Search Queries Used",
        "",
    ]
    lines.extend(f"- {query}" for query in task.search_queries)
    lines.extend(["", "## Top-Ranked Systems Integrators", ""])
    for row in rows:
        company = row.company
        lines.extend(
            [
                f"### {row.rank or '-'}: {company.name}",
                "",
                f"- Score: {row.relevance_score}",
                f"- Website: {company.website or 'Unknown'}",
                f"- Headquarters: {company.headquarters or 'Unknown'}",
                f"- Services: {', '.join(company.services) or 'Unknown'}",
                f"- Vendor partnerships: {', '.join(company.vendor_partnerships) or 'Unknown'}",
                f"- Source URLs: {', '.join(source.source_url for source in company.sources) or 'None'}",
                f"- Scoring breakdown: {json.dumps(row.scoring_breakdown, sort_keys=True)}",
                f"- Notes: {row.notes or company.notes or 'None'}",
                "",
            ]
        )
    lines.extend(
        [
            "## Caveats",
            "",
            "- Configure an official search provider before treating results as production web research.",
            "- Future providers must respect robots.txt, rate limits, and avoid login-only or paywalled data.",
            "- Low-confidence results should be reviewed before outreach.",
            "",
        ]
    )
    Path(path).write_text("\n".join(lines), encoding="utf-8")


def _serialize_result(row: ResearchResult) -> dict[str, object]:
    company = row.company
    return {
        "rank": row.rank,
        "relevance_score": row.relevance_score,
        "scoring_breakdown": row.scoring_breakdown,
        "company": {
            "id": company.id,
            "name": company.name,
            "website": company.website,
            "headquarters": company.headquarters,
            "countries_served": company.countries_served,
            "services": company.services,
            "vendor_partnerships": company.vendor_partnerships,
            "industries_served": company.industries_served,
            "description": company.description,
            "contacts": [
                {
                    "contact_page_url": contact.contact_page_url,
                    "email": contact.email,
                    "phone": contact.phone,
                }
                for contact in company.contacts
            ],
            "sources": [source.source_url for source in company.sources],
        },
        "notes": row.notes,
    }


def _store_docx(session: Session, task: ResearchTask, rows: list[ResearchResult]) -> None:
    document = Document()
    document.add_heading(f"SI Research Report - Task {task.id}", level=0)
    document.add_paragraph(f"Target industry: {task.target_industry}")
    document.add_paragraph(f"Target geography: {task.target_geography}")
    document.add_paragraph(f"Service categories: {', '.join(task.service_categories) or 'Any'}")
    document.add_paragraph(f"Maximum results: {task.max_results}")

    document.add_heading("Search Queries Used", level=1)
    for query in task.search_queries:
        document.add_paragraph(query, style="List Bullet")

    document.add_heading("Top-Ranked Systems Integrators", level=1)
    table = document.add_table(rows=1, cols=6)
    header = table.rows[0].cells
    header[0].text = "Rank"
    header[1].text = "Company"
    header[2].text = "Score"
    header[3].text = "Website"
    header[4].text = "Services"
    header[5].text = "Sources"
    for row in rows:
        company = row.company
        cells = table.add_row().cells
        cells[0].text = str(row.rank or "")
        cells[1].text = company.name
        cells[2].text = f"{row.relevance_score:.1f}"
        cells[3].text = company.website or "Unknown"
        cells[4].text = ", ".join(company.services) or "Unknown"
        cells[5].text = "\n".join(source.source_url for source in company.sources) or "None"

    document.add_heading("Caveats", level=1)
    document.add_paragraph("Review low-confidence results before outreach.", style="List Bullet")
    document.add_paragraph("Only public business contact information should be used.", style="List Bullet")
    document.add_paragraph("Source URLs are included for auditability.", style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    artifact = ReportArtifact(
        task_id=task.id,
        format="docx",
        filename=f"si-research-task-{task.id}.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=buffer.getvalue(),
    )
    session.add(artifact)
